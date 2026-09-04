from __future__ import annotations

import shutil
import zipfile
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parents[1]
REFERENCE = Path(r"C:\Java proyectos\DEA_URBELIX.docx")
OUTPUT = ROOT / "DEA_URBELIX.docx"
WORK_DIR = ROOT / "tmp" / "dea_build"
REFERENCE_WORK = WORK_DIR / "reference_work.docx"
DIAGRAM_DIR = WORK_DIR / "diagrams"

FONT_DIR = Path(r"C:\Windows\Fonts")
FONT_REGULAR = FONT_DIR / "segoeui.ttf"
FONT_BOLD = FONT_DIR / "segoeuib.ttf"

NAVY = "1F4E78"
DEEP_BLUE = "0F4C5C"
BLUE = "2E74B5"
GOLD = "C7952D"
INK = "1F2933"
MUTED = "5E6B75"
GRID = "C7D1D8"
LIGHT_BLUE = "EAF2F8"
LIGHT_TEAL = "E9F4F3"
LIGHT_GOLD = "FFF7E5"
LIGHT_GRAY = "F4F6F7"
GREEN = "1F7A55"
AMBER = "9A6500"
RED = "A12828"


def hex_rgb(value: str) -> RGBColor:
    return RGBColor.from_string(value)


def font_for(size: int, bold: bool = False):
    path = FONT_BOLD if bold and FONT_BOLD.exists() else FONT_REGULAR
    try:
        return ImageFont.truetype(str(path), size)
    except OSError:
        return ImageFont.load_default()


def normalize_reference() -> None:
    WORK_DIR.mkdir(parents=True, exist_ok=True)
    with zipfile.ZipFile(REFERENCE, "r") as source, zipfile.ZipFile(
        REFERENCE_WORK, "w", zipfile.ZIP_DEFLATED
    ) as target:
        for item in source.infolist():
            data = source.read(item.filename)
            if item.filename == "word/document.xml":
                text = data.decode("utf-8")
                text = text.replace('w:left="1842.5196850393704"', 'w:left="1800"')
                data = text.encode("utf-8")
            target.writestr(item, data)


def set_run_font(run, size=None, color=INK, bold=None, italic=None, name="Arial"):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = hex_rgb(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_style_font(style, name="Arial", size=10.5, color=INK, bold=False, italic=False):
    style.font.name = name
    style._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    style._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    style._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), name)
    style.font.size = Pt(size)
    style.font.color.rgb = hex_rgb(color)
    style.font.bold = bold
    style.font.italic = italic


def set_cell_shading(cell, fill: str):
    properties = cell._tc.get_or_add_tcPr()
    shading = properties.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        properties.append(shading)
    shading.set(qn("w:fill"), fill)
    shading.set(qn("w:val"), "clear")


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    properties = cell._tc.get_or_add_tcPr()
    margins = properties.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        properties.append(margins)
    for tag, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = margins.find(qn(f"w:{tag}"))
        if node is None:
            node = OxmlElement(f"w:{tag}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa: list[int], indent=120):
    table.autofit = False
    table.allow_autofit = False
    total = sum(widths_dxa)
    properties = table._tbl.tblPr
    width = properties.first_child_found_in("w:tblW")
    if width is None:
        width = OxmlElement("w:tblW")
        properties.insert(0, width)
    width.set(qn("w:w"), str(total))
    width.set(qn("w:type"), "dxa")
    table_indent = properties.first_child_found_in("w:tblInd")
    if table_indent is None:
        table_indent = OxmlElement("w:tblInd")
        properties.append(table_indent)
    table_indent.set(qn("w:w"), str(indent))
    table_indent.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width_dxa in widths_dxa:
        column = OxmlElement("w:gridCol")
        column.set(qn("w:w"), str(width_dxa))
        grid.append(column)
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            cell.width = Inches(widths_dxa[index] / 1440)
            tc_width = cell._tc.get_or_add_tcPr().first_child_found_in("w:tcW")
            if tc_width is None:
                tc_width = OxmlElement("w:tcW")
                cell._tc.get_or_add_tcPr().append(tc_width)
            tc_width.set(qn("w:w"), str(widths_dxa[index]))
            tc_width.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def mark_header_row(row):
    tr_properties = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_properties.append(header)


def set_paragraph_border(paragraph, color=GOLD, size="12", space="1"):
    properties = paragraph._p.get_or_add_pPr()
    borders = properties.find(qn("w:pBdr"))
    if borders is None:
        borders = OxmlElement("w:pBdr")
        properties.append(borders)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), space)
    bottom.set(qn("w:color"), color)
    borders.append(bottom)


def add_page_field(paragraph):
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run = paragraph.add_run()
    run._r.append(begin)
    run._r.append(instruction)
    run._r.append(separate)
    run._r.append(text)
    run._r.append(end)
    set_run_font(run, size=8.5, color=MUTED)


def set_repeat_table_headers(table):
    if table.rows:
        mark_header_row(table.rows[0])


def set_table_borders(table, color=GRID, size="4"):
    properties = table._tbl.tblPr
    borders = properties.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        properties.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = borders.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), size)
        node.set(qn("w:space"), "0")
        node.set(qn("w:color"), color)


def ensure_numbering(doc, num_fmt, lvl_text):
    numbering = doc.part.numbering_part.element
    existing = [int(node.get(qn("w:numId"))) for node in numbering.findall(qn("w:num"))]
    abstract_ids = [int(node.get(qn("w:abstractNumId"))) for node in numbering.findall(qn("w:abstractNum"))]
    next_num_id = max(existing or [0]) + 1
    next_abstract_id = max(abstract_ids or [0]) + 1
    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(next_abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    level = OxmlElement("w:lvl")
    level.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    level.append(start)
    fmt = OxmlElement("w:numFmt")
    fmt.set(qn("w:val"), num_fmt)
    level.append(fmt)
    text = OxmlElement("w:lvlText")
    text.set(qn("w:val"), lvl_text)
    level.append(text)
    level_properties = OxmlElement("w:pPr")
    indent = OxmlElement("w:ind")
    indent.set(qn("w:left"), "540")
    indent.set(qn("w:hanging"), "270")
    level_properties.append(indent)
    level.append(level_properties)
    abstract.append(level)
    numbering.append(abstract)
    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(next_num_id))
    abstract_reference = OxmlElement("w:abstractNumId")
    abstract_reference.set(qn("w:val"), str(next_abstract_id))
    num.append(abstract_reference)
    numbering.append(num)
    return next_num_id


def add_table(doc, headers, rows, widths, header_fill=NAVY, body_fill=None, first_col_fill=None, font_size=9):
    table = doc.add_table(rows=1, cols=len(headers))
    set_table_geometry(table, widths)
    header = table.rows[0]
    for index, value in enumerate(headers):
        cell = header.cells[index]
        set_cell_shading(cell, header_fill)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(str(value))
        set_run_font(run, size=9, color="FFFFFF", bold=True)
    for row_index, values in enumerate(rows):
        cells = table.add_row().cells
        for index, value in enumerate(values):
            if body_fill:
                set_cell_shading(cells[index], body_fill)
            elif index % 2 == 1:
                set_cell_shading(cells[index], LIGHT_GRAY)
            if first_col_fill and index == 0:
                set_cell_shading(cells[index], first_col_fill)
            p = cells[index].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            run = p.add_run(str(value))
            set_run_font(run, size=font_size, color=INK, bold=(index == 0 and first_col_fill is not None))
    set_repeat_table_headers(table)
    set_table_borders(table)
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.line_spacing = 1.05
    return table


def add_callout(doc, label, text, fill=LIGHT_BLUE, accent=BLUE):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [9360], indent=120)
    set_repeat_table_headers(table)
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.15
    label_run = p.add_run(f"{label}. ")
    set_run_font(label_run, size=10, color=accent, bold=True)
    text_run = p.add_run(text)
    set_run_font(text_run, size=10, color=INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return table


def add_body(doc, text, style="Normal", after=6, keep=False):
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.15
    if keep:
        p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    if style == "Normal":
        set_run_font(run, size=10.5, color=INK)
    return p


def add_bullet(doc, text, numbered=False):
    style = "List Number" if numbered else "List Bullet"
    p = doc.add_paragraph(style=style)
    p.paragraph_format.left_indent = Inches(0.38)
    p.paragraph_format.first_line_indent = Inches(-0.19)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.15
    num_pr = p._p.get_or_add_pPr().find(qn("w:numPr"))
    if num_pr is None:
        num_pr = OxmlElement("w:numPr")
        p._p.get_or_add_pPr().append(num_pr)
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num_id = OxmlElement("w:numId")
    num_id.set(qn("w:val"), str(LIST_NUMBER_ID if numbered else LIST_BULLET_ID))
    num_pr.append(ilvl)
    num_pr.append(num_id)
    run = p.add_run(text)
    set_run_font(run, size=10.2, color=INK)
    return p


def add_caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(9)
    run = p.add_run(text)
    set_run_font(run, size=9, color=MUTED, italic=True)
    return p


def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    p.paragraph_format.keep_with_next = True
    return p


def add_image_alt(inline_shape, description):
    doc_pr = inline_shape._inline.docPr
    doc_pr.set("descr", description)
    doc_pr.set("title", description[:80])


def diagram_canvas(title, subtitle, size=(1600, 900)):
    image = Image.new("RGB", size, "white")
    draw = ImageDraw.Draw(image)
    draw.rectangle((0, 0, size[0], 116), fill="#F4F7F9")
    draw.rectangle((0, 0, 18, size[1]), fill=f"#{NAVY}")
    draw.text((55, 28), title, font=font_for(34, True), fill=f"#{NAVY}")
    draw.text((55, 76), subtitle, font=font_for(20), fill=f"#{MUTED}")
    return image, draw


def box(draw, x, y, w, h, title, body="", fill="#EAF2F8", outline=f"#{NAVY}", title_fill=None):
    draw.rounded_rectangle((x, y, x + w, y + h), radius=18, fill=fill, outline=outline, width=3)
    draw.text((x + 22, y + 18), title, font=font_for(24, True), fill=title_fill or f"#{NAVY}")
    if body:
        lines = body.split("\n")
        for i, line in enumerate(lines):
            draw.text((x + 22, y + 60 + i * 28), line, font=font_for(19), fill=f"#{INK}")


def arrow(draw, start, end, fill=f"#{GOLD}", width=5):
    draw.line((start[0], start[1], end[0], end[1]), fill=fill, width=width)
    import math

    angle = math.atan2(end[1] - start[1], end[0] - start[0])
    length = 18
    left = (end[0] - length * math.cos(angle - 0.45), end[1] - length * math.sin(angle - 0.45))
    right = (end[0] - length * math.cos(angle + 0.45), end[1] - length * math.sin(angle + 0.45))
    draw.polygon([end, left, right], fill=fill)


def make_context_diagram(path: Path):
    image, draw = diagram_canvas("Vista de contexto", "Actores, sistema central e integraciones controladas")
    box(draw, 615, 285, 370, 240, "URBELIX", "Aplicacion web\nGestion residencial", fill=f"#{LIGHT_TEAL}", outline=f"#{DEEP_BLUE}")
    box(draw, 70, 190, 360, 125, "ADMIN", "Gestion y supervision", fill=f"#{LIGHT_BLUE}")
    box(draw, 70, 385, 360, 125, "RESIDENTE", "Autogestion y solicitudes", fill=f"#{LIGHT_BLUE}")
    box(draw, 70, 580, 360, 125, "PORTERIA", "Operacion de accesos", fill=f"#{LIGHT_BLUE}")
    box(draw, 1130, 130, 370, 125, "MySQL", "Produccion + Flyway", fill="#EEF1F3", outline=f"#{MUTED}")
    box(draw, 1130, 300, 370, 125, "SMTP", "Recuperacion y avisos", fill=f"#{LIGHT_GOLD}", outline=f"#{GOLD}")
    box(draw, 1130, 470, 370, 125, "Sandbox local", "Pagos simulados", fill=f"#{LIGHT_GOLD}", outline=f"#{GOLD}")
    box(draw, 1130, 640, 370, 125, "FastAPI opcional", "Reportes PDF", fill=f"#{LIGHT_BLUE}")
    for y in (250, 445, 640):
        arrow(draw, (430, y), (615, 405), fill=f"#{GOLD}")
    for y in (192, 362, 532, 702):
        arrow(draw, (985, 405), (1130, y), fill=f"#{GOLD}")
    image.save(path)


def make_layers_diagram(path: Path):
    image, draw = diagram_canvas("Vista lógica y arquitectura en capas", "Flujo de una solicitud desde el navegador hasta la persistencia")
    layers = [
        ("Navegador", "Thymeleaf + CSS + JS", LIGHT_BLUE),
        ("Seguridad", "Spring Security + CSRF\nroles y sesiones", LIGHT_GOLD),
        ("Controladores", "HTTP + validacion\nrespuestas y vistas", LIGHT_TEAL),
        ("Servicios", "Reglas de negocio\nestados y autorizacion", LIGHT_TEAL),
        ("Persistencia", "JPA + Repositories\ntransacciones", LIGHT_BLUE),
        ("Datos", "H2 dev | MySQL prod\nFlyway en prod", "EEF1F3"),
    ]
    x, y, w, h, gap = 160, 160, 1280, 88, 32
    for i, (title, body, fill) in enumerate(layers):
        yy = y + i * (h + gap)
        box(draw, x, yy, w, h, title, body, fill=f"#{fill}", outline=f"#{NAVY}")
        if i < len(layers) - 1:
            arrow(draw, (800, yy + h), (800, yy + h + gap), fill=f"#{GOLD}", width=4)
    image.save(path)


def make_process_diagram(path: Path):
    image, draw = diagram_canvas("Vista de procesos", "Patron común para los flujos transaccionales de URBELIX")
    steps = [
        ("1", "Solicitud", "Usuario inicia\nla accion"),
        ("2", "Validacion", "Datos, rol y\nreglas del modulo"),
        ("3", "Servicio", "Regla de negocio\ny transaccion"),
        ("4", "Persistencia", "JPA guarda y\nregistra auditoria"),
        ("5", "Respuesta", "Vista, mensaje,\nPDF o descarga"),
    ]
    start_x, y, w, h, gap = 70, 310, 260, 190, 52
    for index, (number, title, body) in enumerate(steps):
        x = start_x + index * (w + gap)
        fill = f"#{LIGHT_BLUE}" if index != 2 else f"#{LIGHT_TEAL}"
        box(draw, x, y, w, h, f"{number}. {title}", body, fill=fill)
        if index < len(steps) - 1:
            arrow(draw, (x + w, y + h / 2), (x + w + gap, y + h / 2), fill=f"#{GOLD}")
    draw.text((260, 660), "Si una validacion falla, el flujo vuelve a la vista con un mensaje claro y conserva los datos no sensibles cuando corresponde.", font=font_for(21), fill=f"#{MUTED}")
    image.save(path)


def make_deployment_diagram(path: Path):
    image, draw = diagram_canvas("Vista de despliegue", "Entorno actual y destino de despliegue documentado")
    box(draw, 80, 230, 330, 180, "Cliente web", "Navegador de escritorio\nHTTP / HTTPS", fill=f"#{LIGHT_BLUE}")
    box(draw, 520, 180, 470, 280, "Servidor URBELIX", "Spring Boot :8080\nThymeleaf + Spring Security\nServicios y auditoria", fill=f"#{LIGHT_TEAL}", outline=f"#{DEEP_BLUE}")
    box(draw, 1110, 120, 360, 150, "H2 archivo", "Perfil dev\n./data/nexurdb", fill="#EEF1F3", outline=f"#{MUTED}")
    box(draw, 1110, 330, 360, 150, "MySQL 8", "Perfil prod\nFlyway V0.1-V14", fill="#EEF1F3", outline=f"#{MUTED}")
    box(draw, 520, 580, 470, 130, "Proveedor opcional", "FastAPI :8000 para PDF de reportes", fill=f"#{LIGHT_BLUE}")
    arrow(draw, (410, 320), (520, 320), fill=f"#{GOLD}")
    arrow(draw, (990, 270), (1110, 195), fill=f"#{GOLD}")
    arrow(draw, (990, 370), (1110, 405), fill=f"#{GOLD}")
    arrow(draw, (755, 460), (755, 580), fill=f"#{GOLD}")
    draw.text((80, 790), "Estado: dev y Docker verificados localmente. Faltan aceptación, backups y dominio.", font=font_for(20, True), fill=f"#{AMBER}")
    image.save(path)


def make_data_diagram(path: Path):
    image, draw = diagram_canvas("Vista de datos", "Relaciones principales del modelo persistente")
    boxes = {
        "Usuario": (70, 200, 300, 120, LIGHT_BLUE),
        "Residente": (70, 430, 300, 120, LIGHT_TEAL),
        "Apartamento": (480, 200, 320, 120, LIGHT_BLUE),
        "Pago": (930, 130, 300, 120, LIGHT_GOLD),
        "Reserva": (930, 320, 300, 120, LIGHT_TEAL),
        "Visitante": (930, 510, 300, 120, LIGHT_BLUE),
        "Incidencia": (480, 520, 320, 120, LIGHT_TEAL),
        "Vehiculo /\nParqueadero": (1280, 320, 250, 150, "EEF1F3"),
    }
    for title, (x, y, w, h, fill) in boxes.items():
        box(draw, x, y, w, h, title, fill=f"#{fill}")
    links = [
        ((370, 260), (480, 260)),
        ((220, 320), (220, 430)),
        ((800, 260), (930, 190)),
        ((800, 260), (930, 380)),
        ((800, 260), (930, 570)),
        ((370, 490), (480, 570)),
        ((1230, 380), (1280, 395)),
    ]
    for start, end in links:
        arrow(draw, start, end, fill=f"#{GOLD}", width=4)
    draw.text((75, 735), "Las relaciones se protegen en los servicios y en la base de datos; las consultas se filtran según el rol autenticado.", font=font_for(20), fill=f"#{MUTED}")
    image.save(path)


def make_roles_diagram(path: Path):
    image, draw = diagram_canvas("Navegación por rol", "La interfaz muestra solamente las tareas permitidas para cada actor")
    columns = [
        (80, "ADMIN", ["Dashboard global", "Usuarios y residentes", "Apartamentos", "Pagos y reportes", "Reservas e incidencias", "Auditoria"], LIGHT_BLUE),
        (600, "RESIDENTE", ["Mi dashboard", "Mis pagos y facturas", "Mis reservas", "Solicitudes de visitantes", "Mis vehiculos", "Mis incidencias"], LIGHT_TEAL),
        (1120, "PORTERIA", ["Dashboard operativo", "Solicitudes de visitantes", "Entrada y salida", "Vehiculos", "Movimientos", "Historial"], LIGHT_GOLD),
    ]
    for x, title, items, fill in columns:
        box(draw, x, 170, 400, 560, title, fill=f"#{fill}", outline=f"#{NAVY}")
        for index, item in enumerate(items):
            yy = 270 + index * 66
            draw.rounded_rectangle((x + 30, yy, x + 370, yy + 45), radius=10, fill="white", outline=f"#{GRID}", width=2)
            draw.text((x + 50, yy + 10), item, font=font_for(20), fill=f"#{INK}")
    draw.text((230, 785), "El control no se limita a ocultar botones: las rutas y métodos también se protegen en Spring Security.", font=font_for(20, True), fill=f"#{NAVY}")
    image.save(path)


def actor(draw, x, y, label):
    draw.ellipse((x, y, x + 42, y + 42), fill=f"#{LIGHT_BLUE}", outline=f"#{NAVY}", width=3)
    draw.line((x + 21, y + 42, x + 21, y + 100), fill=f"#{NAVY}", width=3)
    draw.line((x + 21, y + 58, x - 4, y + 82), fill=f"#{NAVY}", width=3)
    draw.line((x + 21, y + 58, x + 46, y + 82), fill=f"#{NAVY}", width=3)
    draw.line((x + 21, y + 100, x - 2, y + 132), fill=f"#{NAVY}", width=3)
    draw.line((x + 21, y + 100, x + 44, y + 132), fill=f"#{NAVY}", width=3)
    draw.text((x - 14, y + 145), label, font=font_for(18, True), fill=f"#{NAVY}")


def diamond(draw, x, y, w, h, label):
    points = [(x + w / 2, y), (x + w, y + h / 2), (x + w / 2, y + h), (x, y + h / 2)]
    draw.polygon(points, fill=f"#{LIGHT_GOLD}", outline=f"#{GOLD}")
    draw.text((x + 18, y + h / 2 - 11), label, font=font_for(17, True), fill=f"#{INK}")


def lifeline(draw, x, y1, y2):
    for yy in range(y1, y2, 18):
        draw.line((x, yy, x, min(yy + 9, y2)), fill=f"#{GRID}", width=2)


def make_use_case_diagram(path: Path):
    image, draw = diagram_canvas("Diagrama de casos de uso", "Capacidades del sistema organizadas por responsabilidad")
    actor(draw, 100, 190, "ADMIN")
    actor(draw, 100, 430, "RESIDENTE")
    actor(draw, 100, 670, "PORTERIA")
    groups = [
        (430, 155, "Administracion", ["Usuarios y residentes", "Apartamentos", "Pagos y reportes", "Incidencias y avisos"]),
        (890, 155, "Autogestion", ["Mis pagos y factura", "Reservas", "Visitantes", "Vehiculos y PQRS"]),
        (1250, 155, "Operacion", ["Solicitudes", "Entrada y salida", "Parqueaderos", "Historial"]),
    ]
    for x, y, title, items in groups:
        draw.rounded_rectangle((x, y, x + 300, y + 580), radius=20, fill="#FAFCFD", outline=f"#{GRID}", width=3)
        draw.text((x + 24, y + 22), title, font=font_for(22, True), fill=f"#{DEEP_BLUE}")
        for index, item in enumerate(items):
            yy = y + 90 + index * 105
            draw.ellipse((x + 25, yy, x + 275, yy + 52), fill=f"#{LIGHT_TEAL}", outline=f"#{DEEP_BLUE}", width=2)
            draw.text((x + 42, yy + 14), item, font=font_for(16), fill=f"#{INK}")
    for end in ((430, 300), (430, 400)):
        arrow(draw, (150, 300), end, fill=f"#{GOLD}", width=3)
    arrow(draw, (150, 540), (890, 300), fill=f"#{GOLD}", width=3)
    arrow(draw, (150, 780), (1250, 300), fill=f"#{GOLD}", width=3)
    draw.text((390, 775), "Un actor solo accede a los casos que corresponden a su rol.", font=font_for(20, True), fill=f"#{AMBER}")
    image.save(path)


def make_sequence_diagram(path: Path):
    image, draw = diagram_canvas("Diagrama de secuencia: pago y factura", "Flujo de residente, seguridad, servicio, persistencia y respuesta")
    participants = [(90, "Residente"), (365, "Navegador"), (650, "Controller"), (935, "PagoService"), (1220, "Repository"), (1450, "BD")]
    for x, label in participants:
        draw.rounded_rectangle((x - 85, 145, x + 85, 205), radius=12, fill=f"#{LIGHT_BLUE}", outline=f"#{NAVY}", width=3)
        draw.text((x - 65, 165), label, font=font_for(17, True), fill=f"#{NAVY}")
        lifeline(draw, x, 220, 820)
    messages = [
        (260, 90, 365, "1. Solicita pagar"),
        (330, 365, 650, "2. POST /pagos/iniciar"),
        (400, 650, 935, "3. Autoriza y valida"),
        (470, 935, 1220, "4. Crea referencia"),
        (540, 1220, 1450, "5. Guarda PENDIENTE"),
        (610, 935, 650, "6. Retorna checkout"),
        (680, 650, 365, "7. Muestra resultado"),
        (750, 365, 90, "8. Descarga factura"),
    ]
    for y, start, end, label in messages:
        arrow(draw, (start, y), (end, y), fill=f"#{GOLD}", width=3)
        draw.text((min(start, end) + 8, y - 24), label, font=font_for(15), fill=f"#{INK}")
    image.save(path)


def make_activity_diagram(path: Path):
    image, draw = diagram_canvas("Diagrama de actividades: visitante", "Solicitud del residente y decisión operativa de porteria")
    steps = [
        (80, 310, 250, 80, "Residente inicia\nsolicitud"),
        (430, 310, 250, 80, "Sistema valida\ndatos y apartamento"),
        (820, 275, 180, 150, "Datos\nvalidos?"),
        (1130, 170, 300, 90, "Mostrar error y\nconservar formulario"),
        (1130, 410, 300, 90, "Crear solicitud\nPENDIENTE"),
        (1130, 650, 300, 90, "Porteria decide\naprobar o rechazar"),
    ]
    for x, y, w, h, label in steps:
        if "validos" in label:
            diamond(draw, x, y, w, h, label)
        else:
            box(draw, x, y, w, h, label, fill=f"#{LIGHT_TEAL if 'Crear' in label else LIGHT_BLUE}")
    arrow(draw, (330, 350), (430, 350), fill=f"#{GOLD}", width=4)
    arrow(draw, (680, 350), (820, 350), fill=f"#{GOLD}", width=4)
    arrow(draw, (1000, 315), (1130, 215), fill=f"#{RED}", width=4)
    draw.text((1010, 235), "No", font=font_for(18, True), fill=f"#{RED}")
    arrow(draw, (1000, 385), (1130, 455), fill=f"#{GREEN}", width=4)
    draw.text((1010, 425), "Si", font=font_for(18, True), fill=f"#{GREEN}")
    arrow(draw, (1280, 500), (1280, 650), fill=f"#{GOLD}", width=4)
    draw.text((190, 650), "La entrada solo puede registrarse después de una aprobación vigente.", font=font_for(20, True), fill=f"#{AMBER}")
    image.save(path)


def make_state_diagram(path: Path):
    image, draw = diagram_canvas("Diagrama de estados: pago", "Transiciones permitidas y trazabilidad de una obligación")
    states = {
        "PENDIENTE": (100, 340, LIGHT_BLUE), "INICIADO": (420, 180, LIGHT_TEAL),
        "APROBADO": (780, 180, LIGHT_TEAL), "PAGADO": (1130, 180, LIGHT_GOLD),
        "FACTURADO": (1370, 340, LIGHT_GOLD), "RECHAZADO": (420, 550, "FDECEC"),
        "VENCIDO": (780, 650, "FDECEC"), "ANULADO": (1130, 550, "FDECEC"),
    }
    for label, (x, y, fill) in states.items():
        box(draw, x, y, 220, 80, label, fill=f"#{fill}")
    links = [
        ((320, 380), (420, 220), "iniciar"), ((640, 220), (780, 220), "aprobar"),
        ((1000, 220), (1130, 220), "conciliar"), ((1350, 220), (1370, 380), "emitir"),
        ((640, 260), (530, 550), "rechazar"), ((320, 410), (780, 690), "vencer"),
        ((1240, 260), (1240, 550), "anular"),
    ]
    for start, end, label in links:
        arrow(draw, start, end, fill=f"#{GOLD}", width=3)
        draw.text(((start[0] + end[0]) // 2 - 25, (start[1] + end[1]) // 2 - 20), label, font=font_for(15), fill=f"#{INK}")
    draw.text((110, 790), "Un resultado repetido no debe crear un segundo pago: la referencia y la transaccion simulada dejan trazabilidad.", font=font_for(19, True), fill=f"#{AMBER}")
    image.save(path)


def make_class_diagram(path: Path):
    image, draw = diagram_canvas("Diagrama de clases del dominio", "Entidades principales, relaciones y responsabilidades persistentes")
    classes = [
        (70, 150, 300, 210, "Usuario", "id\ncorreo\nrol\nactivo"),
        (470, 150, 300, 210, "Residente", "id\ndocumento\ntelefono\napartamento"),
        (870, 150, 300, 210, "Apartamento", "id\nnumero\ntorre\nestado"),
        (1270, 150, 260, 210, "Pago", "id\nmonto\nfechaPago\nestado"),
        (70, 520, 300, 210, "Reserva", "id\nfecha\nhora\nestado"),
        (470, 520, 300, 210, "Visitante", "id\nnombre\nfechaVisita\nestado"),
        (870, 520, 300, 210, "Incidencia", "id\ntitulo\nprioridad\nestado"),
        (1270, 520, 260, 210, "Vehiculo", "id\nplaca\ntipo\nactivo"),
    ]
    for x, y, w, h, title, body in classes:
        box(draw, x, y, w, h, title, body, fill=f"#{LIGHT_TEAL if title in {'Pago', 'Visitante', 'Incidencia'} else LIGHT_BLUE}")
        draw.line((x, y + 54, x + w, y + 54), fill=f"#{GRID}", width=2)
    relations = [
        ((370, 230), (470, 230), "1"),
        ((770, 230), (870, 230), "1"),
        ((1170, 230), (1270, 230), "1..*"),
        ((220, 360), (220, 520), "1..*"),
        ((620, 360), (620, 520), "1..*"),
        ((1020, 360), (1020, 520), "1..*"),
        ((770, 300), (1270, 590), "0..*"),
    ]
    for start, end, label in relations:
        arrow(draw, start, end, fill=f"#{GOLD}", width=3)
        draw.text(((start[0] + end[0]) // 2 - 10, (start[1] + end[1]) // 2 - 20), label, font=font_for(15, True), fill=f"#{AMBER}")
    image.save(path)


def make_communication_diagram(path: Path):
    image, draw = diagram_canvas("Diagrama de comunicacion: aprobar visitante", "Objetos colaboradores y mensajes numerados del caso de uso")
    nodes = [(130, 310, "Residente"), (500, 150, "VisitanteController"), (900, 310, "VisitanteService"), (1250, 150, "VisitanteRepository"), (1250, 520, "AuditoriaService"), (500, 620, "Porteria")]
    for x, y, label in nodes:
        box(draw, x, y, 250, 90, label, fill=f"#{LIGHT_BLUE if 'Service' not in label else LIGHT_TEAL}")
    messages = [((380, 350), (500, 195), "1 solicitar"), ((750, 195), (900, 350), "2 validar"), ((1150, 350), (1250, 195), "3 guardar"), ((1020, 430), (1250, 565), "4 auditar"), ((900, 400), (625, 620), "5 notificar"), ((500, 650), (380, 395), "6 confirmar")]
    for start, end, label in messages:
        arrow(draw, start, end, fill=f"#{GOLD}", width=4)
        draw.text(((start[0] + end[0]) // 2 - 30, (start[1] + end[1]) // 2 - 20), label, font=font_for(16, True), fill=f"#{INK}")
    image.save(path)


def make_component_diagram(path: Path):
    image, draw = diagram_canvas("Diagrama de componentes", "Dependencias técnicas y límites de integración")
    nodes = [
        (90, 250, 300, 150, "Cliente web", "Thymeleaf\nCSS + JS", LIGHT_BLUE),
        (520, 170, 330, 180, "Web MVC", "Controllers\nSecurity\nValidation", LIGHT_TEAL),
        (520, 470, 330, 180, "Application", "Services\nTransactions\nAudit", LIGHT_TEAL),
        (1010, 170, 300, 150, "Persistence", "JPA\nRepositories", LIGHT_BLUE),
        (1010, 470, 300, 150, "MySQL / H2", "Schema + Flyway", "EEF1F3"),
        (1390, 170, 160, 150, "SMTP", "Correo", LIGHT_GOLD),
        (1390, 470, 160, 150, "Sandbox", "Pagos", LIGHT_GOLD),
    ]
    for x, y, w, h, title, body, fill in nodes:
        box(draw, x, y, w, h, title, body, fill=f"#{fill}")
    for start, end in [((390, 325), (520, 260)), ((685, 350), (685, 470)), ((850, 560), (1010, 545)), ((850, 540), (1390, 245)), ((850, 580), (1390, 545))]:
        arrow(draw, start, end, fill=f"#{GOLD}", width=4)
    draw.text((105, 775), "Los componentes externos son opcionales o simulados; la transaccion local conserva su consistencia.", font=font_for(19, True), fill=f"#{AMBER}")
    image.save(path)


def make_package_diagram(path: Path):
    image, draw = diagram_canvas("Diagrama de paquetes", "Organizacion del codigo para limitar acoplamiento y responsabilidades")
    packages = [
        (100, 180, "controller", "rutas HTTP\nvistas"), (470, 180, "service", "reglas\ntransacciones"),
        (840, 180, "repository", "consultas\nJPA"), (1210, 180, "model", "entidades\nestados"),
        (100, 500, "security", "roles\nCSRF"), (470, 500, "integration", "SMTP\nFastAPI"),
        (840, 500, "templates", "pantallas\nformularios"), (1210, 500, "config", "ambientes\nFlyway"),
    ]
    for x, y, title, body in packages:
        box(draw, x, y, 280, 150, title, body, fill=f"#{LIGHT_TEAL if title in {'service', 'security'} else LIGHT_BLUE}")
    links = [((380, 255), (470, 255)), ((750, 255), (840, 255)), ((1120, 255), (1210, 255)), ((240, 330), (240, 500)), ((610, 330), (610, 500)), ((980, 330), (980, 500)), ((1350, 330), (1350, 500)), ((750, 575), (1210, 575))]
    for start, end in links:
        arrow(draw, start, end, fill=f"#{GOLD}", width=3)
    draw.text((150, 790), "Regla: la vista no accede directamente a la base de datos; el servicio aplica la autorizacion.", font=font_for(19, True), fill=f"#{AMBER}")
    image.save(path)


def prepare_styles(doc):
    normal = doc.styles["Normal"]
    set_style_font(normal, size=10.5)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.15
    for name, size, color, before, after in [
        ("Title", 30, NAVY, 0, 8),
        ("Subtitle", 14, MUTED, 0, 18),
        ("Heading 1", 17, NAVY, 18, 8),
        ("Heading 2", 13, BLUE, 12, 5),
        ("Heading 3", 11.5, DEEP_BLUE, 9, 4),
    ]:
        style = doc.styles[name]
        set_style_font(style, size=size, color=color, bold=name in {"Title", "Heading 1", "Heading 2", "Heading 3"})
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.1
        style.paragraph_format.keep_with_next = True
    global LIST_BULLET_ID, LIST_NUMBER_ID
    LIST_BULLET_ID = ensure_numbering(doc, "bullet", "•")
    LIST_NUMBER_ID = ensure_numbering(doc, "decimal", "%1.")
    style_names = [style.name for style in doc.styles]
    for list_name in ("List Bullet", "List Number"):
        if list_name not in style_names:
            style = doc.styles.add_style(list_name, WD_STYLE_TYPE.PARAGRAPH)
            style.base_style = doc.styles["Normal"]
        else:
            style = doc.styles[list_name]
        set_style_font(style, size=10.2)
        style.paragraph_format.space_after = Pt(3)
        style.paragraph_format.line_spacing = 1.15
    if "Caption URBELIX" not in [style.name for style in doc.styles]:
        caption = doc.styles.add_style("Caption URBELIX", WD_STYLE_TYPE.PARAGRAPH)
    else:
        caption = doc.styles["Caption URBELIX"]
    set_style_font(caption, size=9, color=MUTED, italic=False)


def prepare_section(section):
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.85)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.45)
    section.footer_distance = Inches(0.45)
    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    header.paragraph_format.space_after = Pt(0)
    run = header.add_run("URBELIX | Documento de Especificación de Arquitectura")
    set_run_font(run, size=8.5, color=MUTED)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer.paragraph_format.space_before = Pt(0)
    footer_run = footer.add_run("Página ")
    set_run_font(footer_run, size=8.5, color=MUTED)
    add_page_field(footer)
    set_paragraph_border(header, color=GRID, size="6", space="1")


def add_cover(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(85)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("URBELIX")
    set_run_font(run, size=34, color=NAVY, bold=True)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(5)
    run = p.add_run("DOCUMENTO DE ESPECIFICACIÓN DE ARQUITECTURA")
    set_run_font(run, size=22, color=DEEP_BLUE, bold=True)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(22)
    run = p.add_run("Sistema web de gestión residencial")
    set_run_font(run, size=14, color=MUTED)
    rule = doc.add_paragraph()
    rule.paragraph_format.space_after = Pt(20)
    set_paragraph_border(rule, color=GOLD, size="18", space="1")
    metadata = [
        ["Proyecto", "URBELIX - gestión administrativa, financiera y operativa"],
        ["Carpeta de trabajo", "URBELIXXX"],
        ["Estado", "Documento basado en la implementación actual"],
        ["Versión", "1.0 - 30 de agosto de 2026"],
        ["Equipo", "Equipo de Desarrollo URBELIX"],
    ]
    add_table(doc, ["Dato", "Detalle"], metadata, [2100, 7260], header_fill=NAVY, body_fill=LIGHT_BLUE, font_size=9.5)
    doc.add_paragraph().paragraph_format.space_after = Pt(18)
    add_callout(doc, "Lectura recomendada", "Este DEA distingue lo que ya está implementado de lo que todavía requiere validación operativa. No presenta como productivo ningún flujo que aún dependa de Docker, MySQL de aceptación, SMTP real o proveedores externos.", fill=LIGHT_GOLD, accent=AMBER)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(35)
    run = p.add_run("Ficha técnica: 3114733 | Jornada Noche")
    set_run_font(run, size=10, color=MUTED)
    doc.add_page_break()


def add_document_control(doc):
    add_heading(doc, "HISTORIAL DE REVISIONES", 1)
    add_body(doc, "El historial conserva la evolución del documento y permite relacionar cada entrega con el estado real del código.")
    add_table(doc, ["Fecha", "Versión", "Autor", "Descripción", "Revisado por"], [
        ["30/08/2026", "1.0", "Equipo de Desarrollo URBELIX", "Reestructuración del DEA con base en URBELIXXX, requisitos y evidencias actuales.", "Pendiente de revisión del equipo"],
    ], [1200, 900, 2200, 3400, 1660], header_fill=NAVY, body_fill=LIGHT_BLUE, font_size=8.6)
    add_heading(doc, "Contenido del documento", 1)
    toc = [
        "1. Introducción",
        "2. Generalidades del proyecto",
        "3. Vistas de la arquitectura",
        "3.3 Catálogo completo de casos de uso",
        "3.8 Diagrama de casos de uso",
        "3.9 Diagrama de secuencia",
        "3.10 Diagrama de actividades",
        "3.11 Diagrama de estados",
        "3.12 Diagrama de clases",
        "3.13 Diagrama de comunicación",
        "3.14 Diagrama de componentes",
        "3.15 Diagrama de paquetes",
        "4. Arquitectura en capas y decisiones técnicas",
        "5. Vista de datos",
        "6. Definición de interfaces de usuario",
        "7. Características de calidad",
        "8. Seguridad, privacidad y operación",
        "9. Riesgos, pendientes y trazabilidad",
        "10. Criterios de aceptación arquitectónica",
    ]
    for item in toc:
        add_bullet(doc, item)
    add_callout(doc, "Alcance de lectura", "Las secciones 3 a 6 describen cómo está organizado el sistema. Las secciones 7 a 10 explican cómo se debe comprobar que esa organización se mantiene segura, trazable y preparada para una entrega final.", fill=LIGHT_BLUE, accent=BLUE)
    doc.add_page_break()


def add_introduction(doc):
    add_heading(doc, "1. Introducción", 1)
    add_heading(doc, "1.1 Propósito", 2)
    add_body(doc, "Este documento explica la arquitectura de URBELIX con un lenguaje útil para el equipo que debe mantenerlo, probarlo y presentarlo. La intención no es describir una arquitectura ideal desconectada del código, sino dejar claro qué decisiones se tomaron, qué responsabilidades tiene cada módulo y qué verificaciones faltan para poder declarar una versión final.")
    add_body(doc, "URBELIX funciona como una aplicación web centralizada para conjuntos residenciales. El residente consulta y gestiona su información; administración controla los procesos de negocio; y portería opera exclusivamente los flujos de acceso y parqueaderos. La separación se aplica tanto en la interfaz como en las rutas protegidas del backend.")
    add_heading(doc, "1.2 Alcance", 2)
    add_body(doc, "El alcance actual incluye autenticación, usuarios, residentes, apartamentos, pagos simulados, reservas, visitantes, parqueaderos, vehículos, incidencias/PQRS, avisos, notificaciones, auditoría, reportes y exportaciones. La aplicación principal está implementada en Spring Boot con vistas Thymeleaf y persistencia JPA.")
    add_body(doc, "El perfil de desarrollo utiliza H2 en archivo y datos iniciales controlados. El perfil productivo está preparado para MySQL con migraciones Flyway. Los pagos PSE y tarjeta se ejecutan como sandbox local; no se afirma que exista un cobro real. FastAPI/ReportLab es un proveedor opcional de reportes y existe un generador PDF local de respaldo.")
    add_heading(doc, "1.3 Límites y supuestos", 2)
    for text in [
        "La entrega está orientada a navegador web de escritorio; la compatibilidad móvil no es criterio de aceptación de esta fase.",
        "No se integran talanqueras, cámaras, biometría ni hardware físico del conjunto.",
        "Los secretos de SMTP, MySQL y proveedores externos se cargan por variables de entorno y no forman parte del repositorio.",
        "Docker, carga/estrés, backup/restauración y validación completa de producción son actividades posteriores a la estabilización funcional.",
    ]:
        add_bullet(doc, text)
    add_heading(doc, "1.4 Referencias", 2)
    for text in [
        "REQUISITOS_URBELIX.md, especificación funcional y no funcional consolidada del proyecto.",
        "PROJECT_STATUS.md, estado de avance, evidencias y pendientes.",
        "REQUIREMENTS_TRACEABILITY.md, trazabilidad histórica de requisitos y pruebas.",
        "TESTING_AND_DEPLOYMENT.md y MYSQL_MIGRATION.md, estrategia de pruebas y migración.",
        "Documentación oficial de Spring Boot, Spring Security, JPA/Hibernate, MySQL y Flyway.",
    ]:
        add_bullet(doc, text)
    add_heading(doc, "1.5 Definiciones", 2)
    add_table(doc, ["Término", "Definición"], [
        ["DEA", "Documento de Especificación de Arquitectura."],
        ["H2", "Base de datos embebida utilizada en desarrollo y pruebas rápidas."],
        ["Flyway", "Herramienta que versiona y valida las migraciones de la base productiva."],
        ["PQRS", "Peticiones, quejas, reclamos y sugerencias gestionadas como incidencias."],
        ["Sandbox local", "Flujo de pago simulado que reproduce estados y validaciones sin dinero real."],
    ], [1900, 7460], header_fill=DEEP_BLUE, body_fill=LIGHT_GRAY, font_size=9.1)
    doc.add_page_break()


def add_generalities(doc):
    add_heading(doc, "2. Generalidades del proyecto", 1)
    add_heading(doc, "2.1 Justificación", 2)
    add_body(doc, "La administración de una copropiedad suele repartir su información entre hojas de cálculo, mensajes y registros manuales. Ese modelo dificulta saber quién autorizó un cambio, qué pago corresponde a un apartamento o en qué estado está una solicitud. URBELIX concentra esos procesos en una sola plataforma y busca que cada acción deje una evidencia verificable.")
    add_heading(doc, "2.2 Problema que se resuelve", 2)
    add_body(doc, "El problema principal no es únicamente la falta de una pantalla, sino la falta de un flujo coherente. La administración necesita controlar cartera, reservas e incidencias; el residente necesita autogestionarse sin ver datos de terceros; y portería necesita operar con rapidez sin entrar a módulos administrativos. La arquitectura responde a esa necesidad separando permisos, responsabilidades y datos relacionados.")
    add_heading(doc, "2.3 Objetivos", 2)
    add_table(doc, ["Tipo", "Objetivo"], [
        ["General", "Proveer una aplicación web que centralice la gestión residencial y permita operar con información consistente, trazable y segmentada por rol."],
        ["Específico 1", "Proteger el acceso mediante autenticación, BCrypt, sesiones, CSRF y autorización por rol."],
        ["Específico 2", "Permitir administrar residentes, apartamentos y cuentas con validaciones y relaciones coherentes."],
        ["Específico 3", "Gestionar pagos, fechas, estados, referencias y facturas, manteniendo el sandbox como simulación."],
        ["Específico 4", "Coordinar reservas, visitantes, parqueaderos y PQRS con estados y transiciones explícitas."],
        ["Específico 5", "Generar reportes, exportaciones y auditoría para apoyar decisiones y seguimiento."],
    ], [1700, 7660], header_fill=NAVY, body_fill=LIGHT_BLUE, font_size=9.1)
    add_heading(doc, "2.4 Descripción general del sistema", 2)
    add_body(doc, "La solución es un monolito modular con una entrada web única. Spring Boot recibe la solicitud, Spring Security determina si el usuario puede continuar, el controlador valida y prepara la respuesta, el servicio aplica reglas de negocio y los repositorios consultan o actualizan la base de datos. Las vistas Thymeleaf presentan una interfaz distinta según el rol.")
    add_callout(doc, "Decisión de arquitectura", "Se mantiene una aplicación principal en lugar de separar cada módulo en microservicios. Para el tamaño y la fase actual del proyecto, esta decisión reduce complejidad de despliegue y permite probar las transacciones de forma consistente. FastAPI queda como integración opcional y aislada para reportes.", fill=LIGHT_TEAL, accent=DEEP_BLUE)
    add_heading(doc, "2.5 Stakeholders y responsabilidades", 2)
    add_table(doc, ["Stakeholder", "Responsabilidad en el sistema"], [
        ["ADMIN", "Administra usuarios, residentes, apartamentos, cartera, pagos, reservas, incidencias, avisos, reportes, auditoría y catálogo de parqueaderos. No opera entradas ni salidas."],
        ["RESIDENTE", "Consulta y gestiona sus pagos, reservas, visitantes, vehículos, incidencias y perfil. No accede a datos de terceros ni aprueba accesos."],
        ["PORTERIA", "Aprueba o rechaza solicitudes de visitantes, registra entradas y salidas y opera movimientos de parqueadero. No tiene apartamento, cartera, reservas ni funciones administrativas."],
        ["Equipo de desarrollo", "Mantiene código, pruebas, migraciones, configuración y documentación; debe conservar la separación por capas y roles."],
        ["SENA / evaluadores", "Revisa requisitos, arquitectura, calidad, pruebas y evidencia para la sustentación técnica."],
    ], [2050, 7310], header_fill=DEEP_BLUE, body_fill=LIGHT_GRAY, font_size=8.8)
    doc.add_page_break()


def add_architecture_views(doc, diagrams):
    add_heading(doc, "3. Vistas de la arquitectura", 1)
    add_body(doc, "Las vistas siguientes muestran el mismo sistema desde diferentes preguntas: quién interactúa con él, cómo se mueve una solicitud, dónde viven las reglas, cómo se despliega y qué relaciones sostienen los datos.")
    add_heading(doc, "3.1 Vista de contexto", 2)
    add_body(doc, "El centro de la solución es la aplicación Spring Boot. Los tres actores se conectan al mismo punto de entrada, pero cada uno recibe capacidades diferentes. Las integraciones externas no deben convertirse en una dependencia silenciosa: si SMTP o FastAPI no están disponibles, la aplicación debe conservar la operación local y comunicar la limitación.")
    shape = doc.add_picture(str(diagrams["context"]), width=Inches(6.15))
    add_image_alt(shape, "Diagrama de contexto de URBELIX con actores, aplicación e integraciones")
    add_caption(doc, "Figura 1. Contexto de URBELIX y límites de las integraciones.")
    add_heading(doc, "3.2 Actores y casos de uso principales", 2)
    add_table(doc, ["Actor", "Casos de uso principales", "Resultado esperado"], [
        ["ADMIN", "Gestionar usuarios, apartamentos, pagos, reservas, incidencias, avisos y reportes.", "Información administrativa consistente y auditable."],
        ["RESIDENTE", "Consultar pagos, descargar factura, reservar zonas, solicitar visitantes, gestionar vehículos y PQRS.", "Autogestión limitada a su apartamento y sus registros."],
        ["PORTERIA", "Consultar solicitudes, aprobar/rechazar, registrar entrada/salida y movimientos.", "Control operativo de acceso sin exposición de administración."],
    ], [1500, 4400, 3460], header_fill=NAVY, body_fill=LIGHT_BLUE, font_size=8.8)
    add_heading(doc, "3.3 Catálogo completo de casos de uso", 2)
    add_body(doc, "El catálogo delimita el comportamiento esperado sin confundir una pantalla con un caso de uso. Cada caso identifica el actor autorizado, la acción principal y el resultado que debe poder comprobarse en una prueba.")
    add_table(doc, ["ID", "Actor", "Caso de uso", "Resultado / regla principal"], [
        ["CU-01", "Todos", "Iniciar sesión", "Validar credenciales, rol, sesión y primer ingreso."],
        ["CU-02", "Todos", "Cerrar sesión", "Invalidar la sesión y volver al acceso público."],
        ["CU-03", "Todos", "Recuperar contraseña", "Emitir token temporal y permitir cambio seguro."],
        ["CU-04", "ADMIN", "Gestionar usuarios y roles", "Crear, editar, activar y asignar solo roles válidos."],
        ["CU-05", "ADMIN", "Gestionar residentes", "Relacionar residente con apartamento y validar duplicados."],
        ["CU-06", "ADMIN", "Gestionar apartamentos", "Crear y actualizar unidades sin romper relaciones."],
        ["CU-07", "ADMIN", "Gestionar pagos", "Registrar obligación, fecha, estado, referencia y filtros."],
        ["CU-08", "RESIDENTE", "Pagar obligación simulada", "Recorrer sandbox local sin cobro real."],
        ["CU-09", "RESIDENTE", "Consultar y descargar factura", "Permitir descarga solo de pagos propios."],
        ["CU-10", "ADMIN", "Generar reportes y Excel", "Filtrar información autorizada y exportarla."],
        ["CU-11", "ADMIN/RESIDENTE", "Gestionar reservas", "Crear, revisar, aprobar o cancelar según rol y disponibilidad."],
        ["CU-12", "RESIDENTE", "Solicitar visitante", "Crear solicitud asociada al apartamento y fecha."],
        ["CU-13", "PORTERIA", "Resolver solicitud de visitante", "Aprobar o rechazar y dejar motivo cuando aplique."],
        ["CU-14", "PORTERIA", "Registrar entrada y salida", "Permitir movimientos solo con autorización y estado válido."],
        ["CU-15", "ADMIN/PORTERIA", "Gestionar parqueaderos", "Administrar catálogo o registrar operación según rol."],
        ["CU-16", "RESIDENTE", "Gestionar vehículos propios", "Crear, editar y consultar placas relacionadas."],
        ["CU-17", "RESIDENTE", "Crear incidencia o PQRS", "Registrar asunto, prioridad, descripción y adjuntos."],
        ["CU-18", "ADMIN", "Gestionar incidencia", "Aceptar, rechazar con motivo, resolver y cerrar."],
        ["CU-19", "Todos", "Consultar avisos y notificaciones", "Mostrar información del alcance correspondiente."],
        ["CU-20", "ADMIN", "Consultar auditoría", "Revisar actor, acción, entidad, fecha y detalle."],
    ], [900, 1500, 2600, 4360], header_fill=NAVY, body_fill=LIGHT_BLUE, font_size=7.7)
    add_heading(doc, "3.4 Vista de procesos", 2)
    add_body(doc, "Los flujos de registro, pago, reserva, visitante e incidencia comparten un patrón: primero se identifica al usuario, después se valida el dato y el permiso, luego el servicio aplica la regla y finalmente se guarda el resultado. Las transiciones de estados son parte de la regla de negocio, no una decisión libre de la vista.")
    shape = doc.add_picture(str(diagrams["process"]), width=Inches(6.15))
    add_image_alt(shape, "Proceso común de solicitud, validación, servicio, persistencia y respuesta")
    add_caption(doc, "Figura 2. Flujo común de una operación transaccional.")
    doc.add_page_break()
    add_heading(doc, "3.5 Vista lógica", 2)
    add_body(doc, "La vista lógica hace visible la separación que se debe conservar al cambiar el sistema. La vista no decide si un pago está aprobado ni si un visitante puede entrar; esas decisiones viven en servicios y se ejecutan después de la autorización.")
    shape = doc.add_picture(str(diagrams["layers"]), width=Inches(6.15))
    add_image_alt(shape, "Arquitectura por capas desde el navegador hasta la base de datos")
    add_caption(doc, "Figura 3. Capas principales y dirección del flujo.")
    add_heading(doc, "3.6 Vista de implementación", 2)
    add_table(doc, ["Componente", "Responsabilidad", "Evidencia actual"], [
        ["Controller", "Recibe HTTP, aplica validación de entrada, carga datos para la vista y redirige con mensajes.", "Paquetes `controller`, rutas por módulo y pruebas web."],
        ["Service", "Concentra reglas de estado, autorización por propiedad y transacciones.", "Paquetes `service`, pruebas unitarias de pagos, reservas, visitantes e incidencias."],
        ["Repository", "Abstrae consultas JPA y filtros de persistencia.", "Paquetes `repository` y entidades JPA."],
        ["Model", "Representa entidades, enumeraciones y relaciones del dominio.", "Paquetes `model`, migraciones V0.1-V14."],
        ["Views/static", "Presenta formularios, tablas, mensajes, CSS y JavaScript de interacción.", "`templates`, `static/css`, `static/js`."],
    ], [1750, 4300, 3310], header_fill=DEEP_BLUE, body_fill=LIGHT_GRAY, font_size=8.7)
    add_heading(doc, "3.7 Vista de despliegue", 2)
    add_body(doc, "El entorno de desarrollo usa Spring Boot en el puerto 8080 y H2 en archivo. La composición Docker ya levanta Spring Boot, MySQL 8.4, Flyway y el proveedor opcional de reportes con healthchecks. El destino productivo añade Caddy, HTTPS, límites y persistencia; todavía requiere configuración del dominio y operación de backups.")
    shape = doc.add_picture(str(diagrams["deployment"]), width=Inches(6.15))
    add_image_alt(shape, "Diagrama de despliegue de desarrollo y destino productivo")
    add_caption(doc, "Figura 4. Despliegue local actual y destino documentado.")
    doc.add_page_break()
    add_heading(doc, "3.8 Diagrama de casos de uso", 2)
    add_body(doc, "La vista UML confirma la separación de responsabilidades: ADMIN administra, RESIDENTE se autogestiona y PORTERIA opera accesos. El residente no registra directamente la entrada del visitante; envía una solicitud que portería debe resolver.")
    shape = doc.add_picture(str(diagrams["use_cases"]), width=Inches(6.15))
    add_image_alt(shape, "Diagrama UML de casos de uso de URBELIX por actor")
    add_caption(doc, "Figura 7. Casos de uso agrupados por responsabilidad.")
    add_heading(doc, "3.9 Diagrama de secuencia", 2)
    add_body(doc, "La secuencia representa el pago simulado porque concentra autenticación, validación, servicio, persistencia, referencia y factura. La respuesta no afirma que exista una transacción bancaria real: el flujo es completamente local.")
    shape = doc.add_picture(str(diagrams["sequence"]), width=Inches(6.15))
    add_image_alt(shape, "Diagrama de secuencia del pago simulado y descarga de factura")
    add_caption(doc, "Figura 8. Mensajes del flujo de pago y factura.")
    doc.add_page_break()
    add_heading(doc, "3.10 Diagrama de actividades", 2)
    add_body(doc, "El flujo de visitante muestra las validaciones y decisiones que no deben quedar ocultas en la vista. Si los datos no son válidos, no se crea la solicitud; si son válidos, queda pendiente de decisión operativa.")
    shape = doc.add_picture(str(diagrams["activity"]), width=Inches(6.15))
    add_image_alt(shape, "Diagrama de actividades para solicitar y aprobar un visitante")
    add_caption(doc, "Figura 9. Actividades del caso de uso de visitantes.")
    add_heading(doc, "3.11 Diagrama de estados", 2)
    add_body(doc, "Los estados de pago son explícitos y auditables. Cada transición tiene una causa y evita que una pantalla modifique libremente el estado de una obligación.")
    shape = doc.add_picture(str(diagrams["states"]), width=Inches(6.15))
    add_image_alt(shape, "Diagrama de estados de una obligación de pago")
    add_caption(doc, "Figura 10. Estados y transiciones del pago simulado.")
    doc.add_page_break()
    add_heading(doc, "3.12 Diagrama de clases", 2)
    add_body(doc, "El modelo de clases resume las entidades persistentes que soportan los módulos. Las cardinalidades deben conservarse tanto en las entidades JPA como en las migraciones de base de datos.")
    shape = doc.add_picture(str(diagrams["classes"]), width=Inches(6.15))
    add_image_alt(shape, "Diagrama de clases con entidades principales de URBELIX")
    add_caption(doc, "Figura 11. Clases principales del dominio residencial.")
    add_heading(doc, "3.13 Diagrama de comunicación", 2)
    add_body(doc, "La comunicación muestra los mismos colaboradores desde la perspectiva de los mensajes: controlador, servicio, repositorio, auditoría y portería. Así se puede comprobar quién inicia cada acción y dónde se registra.")
    shape = doc.add_picture(str(diagrams["communication"]), width=Inches(6.15))
    add_image_alt(shape, "Diagrama de comunicación para aprobar un visitante")
    add_caption(doc, "Figura 12. Objetos y mensajes del flujo de visitantes.")
    doc.add_page_break()
    add_heading(doc, "3.14 Diagrama de componentes", 2)
    add_body(doc, "Los componentes se organizan por responsabilidad técnica. SMTP y FastAPI son proveedores opcionales; el sandbox de pagos es local y no reemplaza la consistencia de la aplicación principal.")
    shape = doc.add_picture(str(diagrams["components"]), width=Inches(6.15))
    add_image_alt(shape, "Diagrama de componentes de la aplicación web URBELIX")
    add_caption(doc, "Figura 13. Componentes y dependencias de integración.")
    add_heading(doc, "3.15 Diagrama de paquetes", 2)
    add_body(doc, "La vista de paquetes sirve como regla de organización del código. Los controladores no deben consultar la base directamente, las vistas no deben decidir permisos y el paquete de seguridad debe proteger las rutas además de condicionar la navegación.")
    shape = doc.add_picture(str(diagrams["packages"]), width=Inches(6.15))
    add_image_alt(shape, "Diagrama de paquetes y dependencias internas del proyecto")
    add_caption(doc, "Figura 14. Paquetes principales y dirección de dependencias.")
    doc.add_page_break()


def add_layered_architecture(doc):
    add_heading(doc, "4. Arquitectura en capas y decisiones técnicas", 1)
    add_body(doc, "La arquitectura en capas busca que una modificación se pueda localizar. Si cambia la forma de mostrar una factura, no debería ser necesario cambiar la regla que concilia el pago. Si cambia MySQL por otra instancia compatible, los controladores no deberían conocer ese detalle.")
    add_table(doc, ["Capa", "Tecnología", "Responsabilidad", "Regla de diseño"], [
        ["Presentación", "Thymeleaf, HTML, Bootstrap/CSS y JavaScript", "Formularios, navegación, mensajes y descargas.", "No decide permisos ni estados de negocio."],
        ["Seguridad", "Spring Security, sesiones, BCrypt y CSRF", "Autenticación, autorización y primer ingreso.", "Proteger rutas y métodos; ocultar botones no es suficiente."],
        ["Aplicación", "Controllers y Services", "Orquestar casos de uso y transacciones.", "Validar entrada y delegar la regla al servicio."],
        ["Dominio", "Entidades JPA y enumeraciones", "Representar relaciones, estados y restricciones.", "Evitar estados implícitos y relaciones huérfanas."],
        ["Persistencia", "Spring Data JPA, H2/MySQL y Flyway", "Guardar, consultar y evolucionar el esquema.", "No usar `update` en producción; validar migraciones."],
        ["Integración", "SMTP y FastAPI opcional; sandbox local", "Correo, simulación de pagos y reportes.", "Fallos opcionales no deben borrar ni falsear una transacción local."],
    ], [1500, 2200, 2800, 2860], header_fill=NAVY, body_fill=LIGHT_BLUE, font_size=8.3)
    add_heading(doc, "4.1 Patrones utilizados", 2)
    for text in [
        "MVC con vistas del servidor: separa la representación web de las reglas y la persistencia.",
        "Service Layer: centraliza transiciones de estado, propiedad de registros y operaciones compuestas.",
        "Repository: evita que los controladores construyan consultas directamente.",
        "Filter/Advice transversal: registra auditoría, normaliza errores y aplica reglas de primer ingreso.",
        "Provider con fallback: reportes FastAPI es opcional; el PDF local mantiene la capacidad básica.",
        "Trazabilidad de pagos: el sandbox conserva referencia, resultado, transacción simulada y fecha para evitar ambigüedades.",
    ]:
        add_bullet(doc, text)
    add_heading(doc, "4.2 Decisiones y consecuencias", 2)
    add_table(doc, ["Decisión", "Beneficio", "Costo o riesgo"], [
        ["Monolito modular", "Menos complejidad de despliegue y transacciones sencillas.", "Se debe conservar disciplina para que los módulos no se mezclen."],
        ["H2 en dev", "Arranque rápido y pruebas reproducibles sin depender de MySQL.", "No reemplaza la validación productiva de MySQL/Flyway."],
        ["Pagos simulados", "Permiten demostrar estados y factura sin riesgo financiero.", "No representan una autorización bancaria real."],
        ["PORTERIA como rol técnico separado", "Reduce el riesgo de que una cuenta operativa administre cartera o residentes.", "Requiere mantener rutas y navegación estrictamente separadas."],
    ], [2400, 3300, 3660], header_fill=DEEP_BLUE, body_fill=LIGHT_GRAY, font_size=8.6)
    add_callout(doc, "Criterio de mantenimiento", "Toda nueva funcionalidad debe indicar qué controlador, servicio, entidad, migración, vista y prueba modifica. Si un cambio cruza varias capas, la trazabilidad debe quedar documentada antes de cerrarlo.", fill=LIGHT_GOLD, accent=AMBER)
    doc.add_page_break()


def add_data_view(doc, diagrams):
    add_heading(doc, "5. Vista de datos", 1)
    add_body(doc, "El modelo de datos sostiene tres límites: identidad y permisos, relaciones residenciales y trazabilidad de operaciones. Una acción puede ser válida para un residente y no ser válida para otro; por eso las consultas deben partir del usuario autenticado y no únicamente de un identificador recibido desde el navegador.")
    shape = doc.add_picture(str(diagrams["data"]), width=Inches(6.15))
    add_image_alt(shape, "Relaciones principales del modelo de datos de URBELIX")
    add_caption(doc, "Figura 5. Relaciones principales del dominio.")
    add_heading(doc, "5.1 Entidades y relaciones", 2)
    add_table(doc, ["Grupo", "Entidades", "Relación o regla relevante"], [
        ["Identidad", "Usuario, Rol, PasswordResetToken", "Un usuario tiene un rol y una cuenta puede quedar obligada a cambiar contraseña."],
        ["Residencial", "Apartamento, Residente", "El residente se relaciona con un apartamento para filtrar pagos, reservas y solicitudes."],
        ["Finanzas", "Pago", "El pago conserva fechas, método, estado, referencia, resultado y transacción simulada."],
        ["Reservas", "Reserva", "El rango de tiempo debe ser válido y no cruzarse con otra reserva del espacio."],
        ["Acceso", "Visitante, MovimientoParqueadero, Parqueadero, Vehiculo", "Residente solicita; portería aprueba, registra entrada/salida y movimientos."],
        ["Convivencia", "Incidencia, Comentario, Adjunto, Notificacion", "La incidencia mantiene conversación, evidencias privadas, estados y avisos."],
        ["Trazabilidad", "Auditoria, Aviso, ReporteRegistro", "Las mutaciones autorizadas y la información publicada deben poder revisarse."],
    ], [1750, 3050, 4560], header_fill=NAVY, body_fill=LIGHT_BLUE, font_size=8.5)
    add_heading(doc, "5.2 Persistencia y migraciones", 2)
    add_body(doc, "En desarrollo, Hibernate mantiene el esquema H2 para facilitar el trabajo local. En producción, el perfil activa Flyway, ejecuta migraciones versionadas y deja a Hibernate en modo `validate`. La secuencia V0.1-V14 cubre usuarios, apartamentos, parqueaderos, vehículos, visitantes, incidencias, notificaciones, auditoría, fecha efectiva de pago y trazabilidad del sandbox.")
    add_heading(doc, "5.3 Integridad y consistencia", 2)
    for text in [
        "Las claves foráneas deben impedir pagos, reservas, visitantes o vehículos sin una relación residencial válida.",
        "Los cambios de estado se ejecutan en servicios y deben ser idempotentes cuando el origen sea un evento externo.",
        "La cuenta de portería no debe tener residente ni apartamento asociado.",
        "Los adjuntos de incidencias se almacenan fuera de recursos públicos y se descargan solo con autorización.",
        "La auditoría no incluye contraseñas, secretos ni contenido sensible innecesario.",
    ]:
        add_bullet(doc, text)
    add_heading(doc, "5.4 Protección de la información", 2)
    add_body(doc, "La protección de datos no depende únicamente del modelo: las consultas y descargas deben verificar el rol y la relación del registro con el usuario actual. La factura histórica de un residente, por ejemplo, debe ser descargable solo si el pago le pertenece; ADMIN puede consultar el conjunto conforme a su permiso.")
    doc.add_page_break()


def add_ui_view(doc, diagrams):
    add_heading(doc, "6. Definición de interfaces de usuario", 1)
    add_body(doc, "La interfaz de URBELIX se organiza alrededor de tareas, no de nombres técnicos. Un usuario debería entender qué puede hacer desde el menú y recibir un mensaje claro cuando una regla evita guardar una operación.")
    shape = doc.add_picture(str(diagrams["roles"]), width=Inches(6.15))
    add_image_alt(shape, "Navegación de URBELIX separada para ADMIN, RESIDENTE y PORTERIA")
    add_caption(doc, "Figura 6. Menús y tareas esperadas por rol.")
    add_heading(doc, "6.1 Mapa de navegación", 2)
    add_table(doc, ["Área", "ADMIN", "RESIDENTE", "PORTERIA"], [
        ["Inicio", "Dashboard global", "Dashboard propio", "Dashboard operativo"],
        ["Personas", "Usuarios y residentes", "Perfil propio", "No disponible"],
        ["Residencial", "Apartamentos", "Mi apartamento", "No disponible"],
        ["Finanzas", "Pagos, cuotas, exportación", "Mis pagos, checkout simulado, factura", "No disponible"],
        ["Reservas", "Consulta y aprobación", "Crear y consultar propias", "No disponible"],
        ["Accesos", "No opera portería", "Solicitar visitante y consultar propios", "Aprobar/rechazar, entrada/salida"],
        ["Incidencias", "Gestionar todas", "Crear y consultar propias", "No disponible"],
        ["Reportes", "Generar y descargar", "No disponible", "No disponible"],
    ], [1900, 2490, 2490, 2480], header_fill=NAVY, body_fill=LIGHT_BLUE, font_size=8.4)
    add_heading(doc, "6.2 Principios de interacción", 2)
    for text in [
        "Cada formulario identifica campos obligatorios, formato esperado y motivo de rechazo.",
        "Los mensajes de éxito y error deben aparecer después de la acción y no depender únicamente de una alerta JavaScript.",
        "Las tablas de escritorio mantienen columnas legibles, filtros visibles y acciones ubicadas de forma consistente.",
        "Las descargas muestran qué registro se está descargando: factura, estado de cuenta, reporte o plantilla.",
        "Cuando una acción no pertenece al rol, la opción no se muestra y la ruta responde con acceso denegado.",
    ]:
        add_bullet(doc, text)
    add_heading(doc, "6.3 Estados que el usuario debe entender", 2)
    add_table(doc, ["Módulo", "Estados representativos", "Regla de lectura"], [
        ["Pagos", "PENDIENTE, PAGADO, VENCIDO", "Solo `APPROVED` del sandbox local puede llevar un pago a PAGADO; transferencia y efectivo requieren ADMIN."],
        ["Reservas", "PENDIENTE, APROBADA, RECHAZADA, CANCELADA", "No se aprueba dos veces ni se cruza un espacio ocupado."],
        ["Visitantes", "PENDIENTE, APROBADA, RECHAZADA, DENTRO, FINALIZADA", "La entrada depende de aprobación y la salida de que esté dentro."],
        ["Incidencias", "ABIERTA, EN_REVISION, RESUELTA, CERRADA", "La respuesta y los comentarios mantienen el seguimiento."],
    ], [1600, 3100, 4660], header_fill=DEEP_BLUE, body_fill=LIGHT_GRAY, font_size=8.8)
    doc.add_page_break()


def add_quality(doc):
    add_heading(doc, "7. Características de calidad", 1)
    add_body(doc, "La calidad arquitectónica se evalúa con evidencia, no con una promesa. Por eso cada atributo combina una intención, una decisión técnica y una forma de comprobarlo.")
    add_table(doc, ["Atributo", "Cómo se aborda", "Estado de verificación"], [
        ["Seguridad", "BCrypt, sesiones, CSRF, Spring Security, roles estrictos, almacenamiento privado y secretos fuera del repositorio.", "Implementado en local; falta repetición en producción y revisión de configuración HTTPS."],
        ["Usabilidad", "Vistas por rol, validaciones, mensajes, navegación y tablas de escritorio.", "Parcial; falta checklist visual reproducible en navegador."],
        ["Rendimiento", "Consultas por rol, filtros y separación de proveedores de reportes.", "Pendiente medir tiempos y 100 usuarios concurrentes."],
        ["Confiabilidad", "Transacciones, estados, trazabilidad de pagos y fallback de reportes.", "Parcial; falta backup/restauración y despliegue controlado."],
        ["Mantenibilidad", "Capas, pruebas, migraciones, documentación y auditoría.", "Implementado; debe mantenerse en cada cambio."],
        ["Interoperabilidad", "SMTP y FastAPI mediante configuración externa; pagos sin proveedor externo.", "Parcial; SMTP requiere validación operativa."],
    ], [1700, 5000, 2660], header_fill=NAVY, body_fill=LIGHT_BLUE, font_size=8.6)
    add_heading(doc, "7.1 Rendimiento y capacidad", 2)
    add_body(doc, "Como objetivo de aceptación, las pantallas principales deben responder en menos de tres segundos en uso normal local y el entorno de aceptación debe soportar al menos 100 usuarios concurrentes sin errores críticos. Estos valores son criterios de prueba, no mediciones ya demostradas por la suite actual.")
    add_heading(doc, "7.2 Disponibilidad y recuperación", 2)
    add_body(doc, "El despliegue final debe contar con una estrategia de respaldo y restauración para MySQL, un usuario de base de datos con privilegios mínimos, healthchecks y un procedimiento para recuperar el servicio. Estos elementos aún forman parte del trabajo pendiente de operación.")
    add_heading(doc, "7.3 Pruebas y evidencia", 2)
    add_table(doc, ["Nivel", "Qué cubre hoy", "Qué falta completar"], [
        ["Unitarias", "Servicios, validaciones, estados, pagos simulados, PDF e importación Excel.", "Ampliar casos límite según cada requisito."],
        ["Integración", "Spring Boot, H2, rutas, CSRF, roles, persistencia y vistas.", "Repetir con MySQL/Flyway de aceptación."],
        ["Navegador", "Smoke HTTP y validaciones locales.", "Ejecutar flujos completos con los tres roles en navegador."],
        ["Carga/estrés", "No ejecutadas como criterio de release.", "Definir herramienta, datos, umbrales y reporte reproducible."],
    ], [1550, 3900, 3910], header_fill=DEEP_BLUE, body_fill=LIGHT_GRAY, font_size=8.7)
    add_heading(doc, "7.4 Estándares y buenas prácticas", 2)
    for text in [
        "OWASP: validar entradas, proteger sesiones, evitar exposición de secretos y verificar autorización en servidor.",
        "ISO/IEC 25010 como marco de conversación para funcionalidad, usabilidad, eficiencia, seguridad y mantenibilidad.",
        "Ley 1581 de 2012 y normas colombianas aplicables al tratamiento de datos personales.",
        "Convenciones de Git: cada entrega debe quedar identificada por rama, commit y resultado de pruebas.",
    ]:
        add_bullet(doc, text)
    doc.add_page_break()


def add_security_operation(doc):
    add_heading(doc, "8. Seguridad, privacidad y operación", 1)
    add_heading(doc, "8.1 Autenticación y autorización", 2)
    add_body(doc, "El correo se normaliza antes de autenticar, las contraseñas se almacenan con BCrypt y la sesión se invalida al cerrar sesión. Las rutas privadas se protegen por rol y las operaciones sensibles agregan autorización en el servicio. El primer ingreso puede obligar a cambiar la contraseña antes de continuar al resto de la aplicación.")
    add_heading(doc, "8.2 Separación de roles", 2)
    add_table(doc, ["Rol", "Puede hacer", "No puede hacer"], [
        ["ADMIN", "Administrar información y decisiones del conjunto.", "Operar entradas, salidas y movimientos físicos de portería."],
        ["RESIDENTE", "Gestionar registros propios y crear solicitudes.", "Consultar terceros, aprobar visitantes o entrar a rutas administrativas."],
        ["PORTERIA", "Operar visitantes y parqueaderos.", "Ser residente, tener apartamento, gestionar cartera, reservas o módulos administrativos."],
    ], [1550, 3900, 3910], header_fill=NAVY, body_fill=LIGHT_BLUE, font_size=8.8)
    add_heading(doc, "8.3 Datos sensibles y archivos", 2)
    for text in [
        "No se guardan contraseñas, claves privadas ni secretos de proveedores en el código.",
        "Los tokens de recuperación expiran y se invalidan después de usarse.",
        "Los adjuntos de PQRS se almacenan fuera de recursos públicos, con extensión/tipo/tamaño validados.",
        "Los reportes y facturas deben respetar el alcance del rol y no incluir pagos o datos de otro apartamento.",
        "La auditoría registra actor, acción, entidad, fecha y detalle no sensible para reconstruir una mutación.",
    ]:
        add_bullet(doc, text)
    add_heading(doc, "8.4 Configuración por ambiente", 2)
    add_table(doc, ["Aspecto", "Dev", "Prod / aceptación"], [
        ["Base de datos", "H2 en archivo `./data/nexurdb`.", "MySQL 8 con Flyway y `validate`."],
        ["Datos iniciales", "Seed controlado de admin, portería, apartamentos y parqueaderos.", "No se ejecuta seed automático."],
        ["Pagos", "Sandbox local activo por defecto.", "No se realizan cobros ni llamadas a proveedores externos."],
        ["Reportes", "Generador local; FastAPI opcional.", "Configurar proveedor y healthcheck si se usa."],
        ["Correo", "Desactivado salvo variables explícitas.", "SMTP con contraseña de aplicación fuera de Git."],
    ], [1900, 3500, 3960], header_fill=DEEP_BLUE, body_fill=LIGHT_GRAY, font_size=8.7)
    add_callout(doc, "Regla operativa", "Nunca se debe usar la cuenta root de MySQL como usuario permanente de la aplicación. La validación del entorno de aceptación debe usar una cuenta con permisos mínimos y un respaldo previo.", fill=LIGHT_GOLD, accent=AMBER)
    doc.add_page_break()


def add_risks_traceability(doc):
    add_heading(doc, "9. Riesgos, pendientes y trazabilidad", 1)
    add_body(doc, "El proyecto tiene una base funcional amplia, pero todavía hay diferencias entre una demostración local y una versión lista para desplegar. Esta sección deja visibles esas diferencias para que no se pierdan al preparar las pruebas.")
    add_heading(doc, "9.1 Riesgos principales", 2)
    add_table(doc, ["Riesgo", "Impacto", "Mitigación propuesta", "Estado"], [
        ["Migraciones o consultas distintas en MySQL", "Arranque fallido o datos inconsistentes.", "Respaldar, ejecutar Flyway V0.1-V14 y repetir smoke contra MySQL.", "Pendiente"],
        ["SMTP mal configurado", "Recuperación o aviso no llega al usuario.", "Usar contraseña de aplicación, probar envío real y conservar mensaje seguro.", "Parcial"],
        ["Permiso excesivo por ruta nueva", "Un rol ve o modifica información ajena.", "Pruebas negativas por rol y autorización dentro del servicio.", "Parcial"],
        ["Resultado repetido de pago", "Pago conciliado dos veces o trazabilidad ambigua.", "El servicio solo procesa pagos pendientes/vencidos y persiste un identificador de simulación.", "Implementado en sandbox"],
        ["Falta de backup", "Pérdida de cartera, auditoría o incidencias.", "Crear procedimiento automatizado y prueba de restauración.", "Pendiente"],
        ["Rendimiento sin medición", "Degradación al crecer el conjunto.", "Definir carga, umbrales y reporte de resultados.", "Pendiente"],
    ], [2050, 2300, 3900, 1110], header_fill=NAVY, body_fill=LIGHT_BLUE, font_size=8.1)
    add_heading(doc, "9.2 Pendientes que bloquean la versión final", 2)
    for text in [
        "Repetir pruebas de login, roles, registro, pagos, visitantes e incidencias en navegador.",
        "Validar MySQL/Flyway de aceptación con respaldo y usuario de mínimos privilegios.",
        "Probar SMTP real y documentar el comportamiento cuando el correo falla.",
        "Definir y ejecutar pruebas de carga, estrés, recuperación y restauración.",
        "Completar la revisión visual de las vistas de escritorio y ajustar detalles de UX.",
    ]:
        add_bullet(doc, text)
    add_heading(doc, "9.3 Trazabilidad de arquitectura", 2)
    add_table(doc, ["Necesidad", "Decisión arquitectónica", "Evidencia / prueba"], [
        ["Cada rol ve solo lo que necesita", "Autorización por rol en rutas y servicios; navegación condicionada.", "NexurIntegrationTest, smoke HTTP y pruebas 403."],
        ["El residente consulta solo lo propio", "Relación por usuario/apartamento en servicios y repositorios.", "Pruebas de pagos, reservas, vehículos, visitantes e incidencias."],
        ["Los pagos tienen trazabilidad", "Fecha, referencia, método, estado, resultado y transacción simulada.", "PagoServiceTest, PagoSimulacionServiceTest y factura PDF."],
        ["Los cambios se pueden revisar", "AuditoriaRequestFilter, AuditoriaService y entidad Auditoria.", "AuditoriaController y pruebas de persistencia."],
        ["El sistema sobrevive a fallos opcionales", "PDF local de respaldo y correo sin borrar la operación.", "ReporteControllerTest y pruebas de notificaciones."],
    ], [2450, 4000, 2910], header_fill=DEEP_BLUE, body_fill=LIGHT_GRAY, font_size=8.4)
    add_callout(doc, "Fuente de verdad", "El listado detallado de requisitos está en REQUISITOS_URBELIX.md. Este DEA explica la estructura que permite cumplirlos y debe actualizarse cuando cambien roles, módulos, integraciones o decisiones de despliegue.", fill=LIGHT_TEAL, accent=DEEP_BLUE)
    doc.add_page_break()


def add_acceptance(doc):
    add_heading(doc, "10. Criterios de aceptación arquitectónica", 1)
    add_body(doc, "La arquitectura se puede considerar lista para la siguiente fase cuando los flujos funcionales estén cerrados y exista evidencia repetible de que las capas, los roles, los datos y el despliegue se comportan como se describe aquí.")
    add_heading(doc, "10.1 Flujos prioritarios", 2)
    flows = [
        ["FA-01", "Registro y primer ingreso", "Crear residente con apartamento y código válidos; exigir cambio de contraseña y rechazar duplicados o relaciones inválidas."],
        ["FA-02", "Separación estricta de roles", "ADMIN, RESIDENTE y PORTERIA ven y ejecutan únicamente sus tareas; las rutas prohibidas responden 403."],
        ["FA-03", "Pago simulado y factura", "Crear obligación, iniciar checkout local, recorrer estados, conciliar APPROVED y descargar factura autorizada."],
        ["FA-04", "Visitante y portería", "Residente solicita; portería aprueba/rechaza y registra entrada/salida respetando estados."],
        ["FA-05", "Incidencia con trazabilidad", "Residente crea y consulta; administración gestiona estado, comentarios, evidencias y notificaciones."],
    ]
    add_table(doc, ["ID", "Flujo", "Criterio de aceptación"], flows, [1000, 2500, 5860], header_fill=NAVY, body_fill=LIGHT_BLUE, font_size=8.7)
    add_heading(doc, "10.2 Lista de salida", 2)
    checks = [
        "La suite `mvnw.cmd clean test` termina sin fallos y el resultado queda asociado al commit.",
        "El perfil dev inicia con H2 y los tres roles recorren sus rutas principales.",
        "El perfil prod arranca contra MySQL de aceptación y Flyway valida las migraciones.",
        "No existen rutas nuevas sin autorización ni formularios POST sin CSRF.",
        "La simulación de pagos conserva su etiqueta de sandbox y no realiza cargos reales.",
        "Los reportes y facturas respetan el rol y los filtros seleccionados.",
        "Existe respaldo/restauración probado y cuenta MySQL con mínimos privilegios.",
        "Docker, healthchecks, pruebas de navegador y carga quedan documentados antes del release.",
    ]
    for text in checks:
        add_bullet(doc, text, numbered=True)
    add_heading(doc, "10.3 Cierre", 2)
    add_body(doc, "URBELIX ya cuenta con una base arquitectónica coherente para continuar: la aplicación tiene módulos identificables, reglas en servicios, roles explícitos, persistencia versionada y pruebas sobre los flujos críticos. El paso siguiente no es crear otra estructura, sino cerrar las validaciones operativas que convierten esa base en una entrega reproducible.")
    add_callout(doc, "Estado de este DEA", "Documento preparado para revisión del equipo Scrum y para servir como guía de implementación y pruebas. Debe actualizarse si cambia la decisión de roles, pagos simulados, proveedor de reportes o estrategia de despliegue.", fill=LIGHT_GOLD, accent=AMBER)


def build_document():
    normalize_reference()
    # Start from a clean package so unused media from the reference cannot travel into the deliverable.
    doc = Document()
    body = doc._element.body
    sect_pr = body.sectPr
    for child in list(body):
        if child is not sect_pr:
            body.remove(child)
    prepare_styles(doc)
    for section in doc.sections:
        prepare_section(section)
    doc.core_properties.title = "Documento de Especificación de Arquitectura - URBELIX"
    doc.core_properties.subject = "Arquitectura del sistema web de gestión residencial"
    doc.core_properties.author = "Equipo de Desarrollo URBELIX"
    doc.core_properties.comments = "Generado a partir del estado actual de URBELIXXX"

    DIAGRAM_DIR.mkdir(parents=True, exist_ok=True)
    diagrams = {
        "context": DIAGRAM_DIR / "contexto.png",
        "layers": DIAGRAM_DIR / "capas.png",
        "process": DIAGRAM_DIR / "procesos.png",
        "deployment": DIAGRAM_DIR / "despliegue.png",
        "data": DIAGRAM_DIR / "datos.png",
        "roles": DIAGRAM_DIR / "roles.png",
        "use_cases": DIAGRAM_DIR / "casos_de_uso.png",
        "sequence": DIAGRAM_DIR / "secuencia_pago.png",
        "activity": DIAGRAM_DIR / "actividad_visitante.png",
        "states": DIAGRAM_DIR / "estados_pago.png",
        "classes": DIAGRAM_DIR / "clases_dominio.png",
        "communication": DIAGRAM_DIR / "comunicacion_visitante.png",
        "components": DIAGRAM_DIR / "componentes.png",
        "packages": DIAGRAM_DIR / "paquetes.png",
    }
    make_context_diagram(diagrams["context"])
    make_layers_diagram(diagrams["layers"])
    make_process_diagram(diagrams["process"])
    make_deployment_diagram(diagrams["deployment"])
    make_data_diagram(diagrams["data"])
    make_roles_diagram(diagrams["roles"])
    make_use_case_diagram(diagrams["use_cases"])
    make_sequence_diagram(diagrams["sequence"])
    make_activity_diagram(diagrams["activity"])
    make_state_diagram(diagrams["states"])
    make_class_diagram(diagrams["classes"])
    make_communication_diagram(diagrams["communication"])
    make_component_diagram(diagrams["components"])
    make_package_diagram(diagrams["packages"])

    add_cover(doc)
    add_document_control(doc)
    add_introduction(doc)
    add_generalities(doc)
    add_architecture_views(doc, diagrams)
    add_layered_architecture(doc)
    add_data_view(doc, diagrams)
    add_ui_view(doc, diagrams)
    add_quality(doc)
    add_security_operation(doc)
    add_risks_traceability(doc)
    add_acceptance(doc)

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUTPUT))
    print(f"Created {OUTPUT}")
    print(f"Diagrams: {len(diagrams)}")


if __name__ == "__main__":
    build_document()
